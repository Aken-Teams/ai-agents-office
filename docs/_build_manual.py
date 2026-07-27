# -*- coding: utf-8 -*-
"""Build the Email 資料源・多人同時使用 操作手冊 (.docx). Style ref: 操作手冊.pdf."""
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

CJK  = "Microsoft JhengHei"
TEAL = RGBColor(0x0F, 0x7B, 0x6C)
INK  = RGBColor(0x1F, 0x29, 0x33)
SUB  = RGBColor(0x52, 0x60, 0x6D)
MUTE = RGBColor(0x7B, 0x84, 0x94)
AMBER= RGBColor(0xB2, 0x6A, 0x00)
RED  = RGBColor(0xC0, 0x36, 0x2C)
GREEN= RGBColor(0x0F, 0x8A, 0x5F)
WHITE= RGBColor(0xFF, 0xFF, 0xFF)
CONTENT_W = Cm(16.0)

doc = Document()
st = doc.styles['Normal']
st.font.name = CJK; st.font.size = Pt(11)
st.element.rPr.rFonts.set(qn('w:eastAsia'), CJK)
for s in doc.sections:
    s.top_margin = Cm(2.2); s.bottom_margin = Cm(2.2)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.5)

def _cjk(run):
    run.font.name = CJK
    run._element.get_or_add_rPr().get_or_add_rFonts().set(qn('w:eastAsia'), CJK)

def shade(cell, hexcolor):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:val'),'clear'); shd.set(qn('w:fill'),hexcolor); tcPr.append(shd)

def borders(cell, color="E4E7EB", sz="6", edges=('top','left','bottom','right')):
    tcPr = cell._tc.get_or_add_tcPr(); b = OxmlElement('w:tcBorders')
    for edge in edges:
        e = OxmlElement(f'w:{edge}'); e.set(qn('w:val'),'single'); e.set(qn('w:sz'),sz)
        e.set(qn('w:space'),'0'); e.set(qn('w:color'),color); b.append(e)
    tcPr.append(b)

def set_w(cell, w):
    cell.width = w
    tcPr = cell._tc.get_or_add_tcPr(); tw = OxmlElement('w:tcW')
    tw.set(qn('w:w'), str(int(w.twips))); tw.set(qn('w:type'),'dxa'); tcPr.append(tw)

def para(text="", size=11, color=INK, bold=False, align=None, before=2, after=4, indent=None, lh=1.3):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(before); pf.space_after = Pt(after); pf.line_spacing = lh
    if align: p.alignment = align
    if indent is not None: pf.left_indent = indent
    if text:
        r = p.add_run(text); r.bold = bold; r.font.size = Pt(size); r.font.color.rgb = color; _cjk(r)
    return p

def feature_head(tag, title):
    p = doc.add_paragraph(); pf = p.paragraph_format
    pf.space_before = Pt(18); pf.space_after = Pt(7)
    r = p.add_run(tag + "　"); r.font.size = Pt(15); r.bold = True; r.font.color.rgb = TEAL; _cjk(r)
    r2 = p.add_run(title); r2.font.size = Pt(15); r2.bold = True; r2.font.color.rgb = INK; _cjk(r2)
    pr = p._p.get_or_add_pPr(); pbdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom'); bottom.set(qn('w:val'),'single'); bottom.set(qn('w:sz'),'12')
    bottom.set(qn('w:space'),'6'); bottom.set(qn('w:color'),'0F7B6C'); pbdr.append(bottom); pr.append(pbdr)

def subhead(text):
    p = doc.add_paragraph(); pf = p.paragraph_format; pf.space_before = Pt(12); pf.space_after = Pt(3)
    r = p.add_run("▎"); r.font.size = Pt(11.5); r.font.color.rgb = TEAL; _cjk(r)
    r2 = p.add_run(text); r2.font.size = Pt(12); r2.bold = True; r2.font.color.rgb = INK; _cjk(r2)

def item(label, title, desc, label_color=TEAL):
    p = doc.add_paragraph(); pf = p.paragraph_format; pf.space_before = Pt(8); pf.space_after = Pt(1)
    r1 = p.add_run(label); r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = label_color; _cjk(r1)
    r2 = p.add_run("　"); _cjk(r2)
    r3 = p.add_run(title); r3.bold = True; r3.font.size = Pt(11.5); r3.font.color.rgb = INK; _cjk(r3)
    pd = doc.add_paragraph(); pd.paragraph_format.space_after = Pt(3)
    pd.paragraph_format.left_indent = Cm(0.75); pd.paragraph_format.line_spacing = 1.25
    rd = pd.add_run(desc); rd.font.size = Pt(10.5); rd.font.color.rgb = SUB; _cjk(rd)

def fig_placeholder(caption):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0,0); shade(cell, "F7F9FB"); borders(cell, "D8DEE5", "6"); set_w(cell, CONTENT_W)
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(18)
    r = p.add_run("［ 請在此貼上示意圖 ］"); r.font.size = Pt(10.5); r.font.color.rgb = MUTE; _cjk(r)
    cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4); cap.paragraph_format.space_after = Pt(12)
    rc = cap.add_run(caption); rc.font.size = Pt(9.5); rc.font.color.rgb = MUTE; rc.italic = True; _cjk(rc)

def callout(label, text, fill="E6F4F1", bar="0F7B6C", labelcolor=TEAL):
    t = doc.add_table(rows=1, cols=1); t.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = t.cell(0,0); shade(cell, fill); borders(cell, bar, "6"); set_w(cell, CONTENT_W)
    p = cell.paragraphs[0]; p.paragraph_format.space_before = Pt(7); p.paragraph_format.space_after = Pt(7)
    r1 = p.add_run("  " + label + "　"); r1.bold = True; r1.font.size = Pt(11); r1.font.color.rgb = labelcolor; _cjk(r1)
    r2 = p.add_run(text); r2.font.size = Pt(10.5); r2.font.color.rgb = INK; _cjk(r2)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

# ---------------- Cover ----------------
para("操作手冊", size=30, color=TEAL, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, before=10, after=2)
para("Email 資料源・多人同時使用 — 功能操作說明", size=14, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, after=4)
para("適用對象：一般使用者／管理者　|　版本 v1.0　|　共 2 項功能", size=10.5, color=MUTE,
     align=WD_ALIGN_PARAGRAPH.CENTER, after=6)
para("本手冊說明兩項與信件（Email）相關的更新：一是把 Email 當成 AI 產文件的資料來源，"
     "二是多人同時使用時的穩定性強化。內容以實際操作與畫面為主，方便快速上手。",
     size=10.5, color=SUB, align=WD_ALIGN_PARAGRAPH.CENTER, after=8, lh=1.4)

# ---------------- Feature 1 ----------------
feature_head("功能 ①", "Email 可當 AI 資料源（產文件／AI 團隊）")
para("在「產文件」或「AI 團隊」時勾選「我的信件」，AI 就能讀取你自己的信件內容來產出文件，"
     "例如把一封信整理成 Word 報告、把多封往來信彙整成 PPT。系統只讀取你本人有權限的信件，"
     "且信件的附件文字與內文圖片都會完整讀取，不會任意截斷，確保產出內容準確。",
     size=11, color=SUB, lh=1.4)

fig_placeholder("圖 1：Email 資料源 — 點資料源 → 勾我的信件 → AI 產文件")

subhead("操作步驟")
item("步驟 1", "點「資料源」圖示", "在對話框或 AI 團隊的工具列，點一下資料庫（database）圖示，開啟資料源選單。")
item("步驟 2", "勾選「我的信件」", "勾選後即授權「這一次執行」可讀取你的信箱；可同時勾選其他資料源。")
item("步驟 3", "輸入需求，AI 產文件", "描述你要的文件內容，AI 會讀取信件、整理重點，產出 Word／PPT／Excel 等檔案。")

subhead("安全設計")
para("系統以標準的 MCP 連接器接取信箱，透過四個原則確保安全：本人權杖臨時授權、只讀不改、"
     "用完即收、跨使用者隔離。", size=11, color=SUB, lh=1.4)
fig_placeholder("圖 2：MCP 安全設計 — 本人權杖・只讀・用完即收・跨使用者隔離")
fig_placeholder("圖 3：資料隔離 — A 只讀得到 A 的信，讀不到 B 的信")

callout("核心精神", "資料只走你本人的權限，AI 不會、也無法讀到別人的信箱；連線用完即收。")

# ---------------- Feature 2 ----------------
doc.add_page_break()
feature_head("功能 ②", "多人同時使用更順（限流保護）")
para("上課時常有約 30 人同時開啟信件助手，眾人的背景更新與資料載入會同時湧向信箱 API，"
     "容易造成塞車、甚至被信箱 API 回覆「請求太多（429）」而出錯。本次針對尖峰情境做了強化，"
     "讓多人同時使用也能順暢，不再無限轉圈或跳出錯誤。", size=11, color=SUB, lh=1.4)

fig_placeholder("圖 4：排隊閘門 — 爆量請求排隊逐一送出，不再一次灌爆")

subhead("做了哪些強化")
item("①", "請求排隊", "把同時打信箱 API 的請求量設上限，其餘自動排隊、稍後送出，避免一次灌爆。")
item("②", "自動退避重試", "萬一被回 429，系統會依規範等待後自動重試，多數情況使用者完全無感。")
item("③", "友善載入提示", "尖峰載入較久時，顯示「多人同時使用中，正在為你載入…」，不再像卡死。")

fig_placeholder("圖 5：自動退避重試 — 429 → 等一下 → 重試成功，使用者不會看到錯誤")
fig_placeholder("圖 6：尖峰提示 — 「多人同時使用中，正在為你載入…」，稍等即出信")

subhead("使用者體驗：調整前 vs 調整後")
rows = [
    ("情境", "調整前", "調整後"),
    ("30 人同時開啟", "一直轉圈、偶爾出錯", "自動排隊，稍等即載入"),
    ("被信箱 API 限流", "畫面卡住或跳錯誤", "自動退避重試，多半無感"),
]
tb = doc.add_table(rows=len(rows), cols=3); tb.alignment = WD_TABLE_ALIGNMENT.CENTER; tb.style = 'Table Grid'
widths = [Cm(4.2), Cm(5.9), Cm(5.9)]
for ri, row in enumerate(rows):
    for ci, val in enumerate(row):
        c = tb.cell(ri, ci); set_w(c, widths[ci])
        pp = c.paragraphs[0]; pp.paragraph_format.space_before = Pt(3); pp.paragraph_format.space_after = Pt(3)
        pp.paragraph_format.left_indent = Cm(0.15)
        r = pp.add_run(val); r.font.size = Pt(10.5); _cjk(r)
        if ri == 0:
            r.bold = True; r.font.color.rgb = WHITE; shade(c, "0F7B6C")
        elif ci == 0:
            r.bold = True; r.font.color.rgb = INK; shade(c, "F1F5F4")
        elif ci == 1:
            r.font.color.rgb = RED
        else:
            r.bold = True; r.font.color.rgb = GREEN

subhead("（管理者）如何自行驗證")
para("到後台〔安全審計 → 信件 Gateway 限流閘門〕，點「開始壓測（模擬 N 人同時）」，當場即可看到"
     "「全部成功、失敗到前端＝0」的結果。平時該卡片的「失敗到前端」數字應維持 0。",
     size=11, color=SUB, lh=1.4)
fig_placeholder("圖 7：後台驗證 — 壓測「模擬 30 人同時」全部成功、失敗到前端 = 0")

callout("核心精神", "多人同時使用也穩定；尖峰時只是稍等，不會壞、不會跳錯。")
callout("小提醒", "若上課尖峰仍遇到異常，請截圖回報，我們可依後台壓測數據快速定位原因。",
        fill="FDF3E2", bar="ECCB8F", labelcolor=AMBER)

out = r"D:\github\ai-agents-office\docs\操作手冊-Email資料源與多人同時使用.docx"
doc.save(out)
print("SAVED OK")
